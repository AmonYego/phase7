import streamlit as st
from PyPDF2 import PdfReader
import google.generativeai as genai
from docx import Document
from streamlit import spinner

st.set_page_config(page_title="ClassRoom AI", page_icon="ğŸ“")
st.markdown("""
<div style='display: flex; justify-content: space-between; align-items: center; padding-top: 30px; color: blue; font-size: 20px;'>
    <div style='text-align: left;'><b>Â©Yego Senior</b></div>
    <div style='text-align: center;'><b>ğŸ†Helping you drive good grades homeğŸ“.</b></div>
    <div style='text-align: right;'><b>Â©ClassRoom AIğŸ“</b></div>
</div>
""", unsafe_allow_html=True)


st.markdown("<h1 style='text-align: center; color: blue;'>ClassRoom AIğŸ“</h1>", unsafe_allow_html=True)

st.markdown("<h3 style='text-align: center;'>Ace your exams â€” Revise Like a Pro with AI</h3>", unsafe_allow_html=True)

st.write("Upload your class notes and past papers for AI-powered topic analysis, simple explanations, and smart practice questions â€” get your answered papers marked instantly with AI, or chat directly for explanations, help, and revision support.")
genai.configure(api_key="AIzaSyARKbi8gr-3sLsw5KOEsZMUsudHA53sxBA")
model = genai.GenerativeModel("gemini-2.5-flash")

level = st.selectbox(
    "ğŸ“š Select your education level:",
    ["-- Select Level --", "Lower Primary (Grade 1-5)", "Upper Primary (Grade 6-9)",
     "High School (Grade 10-12)", "College/University"]
)

def get_level_prompt(level):
    if "Lower Primary" in level:
        return """
        The learner is in LOWER PRIMARY (Grade 1â€“5 in Kenyan CBC).
        âœ… Use VERY SIMPLE English or Kiswahili-friendly language.
        âœ… Explain like teaching a young child (short, joyful, friendly sentences).
        âœ… Give fun and relatable examples (toys, animals, food, games, family, cartoons).
        âœ… Use a playful and encouraging tone (e.g. â€œGreat job! Letâ€™s try another one!â€).
        âœ… Avoid complex terms, abstract ideas, or deep reasoning.
        âœ… Encourage learning with excitement and imagination.
        âœ… If explaining a concept:
            1. Say what it is in 1 easy line.
            2. Give one playful example.
            3. Ask a tiny question or give a fun quick activity (optional).
        """

    elif "Upper Primary" in level:
        return """
        The learner is in UPPER PRIMARY (Grade 6â€“9 in Kenyan CBC).
        âœ… Use simple English with mild academic structure.
        âœ… Explain concepts step-by-step with clear logic.
        âœ… Use relatable examples from school, hobbies, friends, or basic science.
        âœ… Introduce basic academic terms and define them simply.
        âœ… Tone should be supportive, like a kind school tutor.
        âœ… If explaining a concept:
            1. Give a short definition.
            2. Explain step-by-step using relatable examples.
            3. Give one worked-out example.
            4. Point out one common mistake or misconception.
            5. Provide a 2-question practice quiz with answers.
            6. End with one â€œchallengeâ€ question for extra thinking.
        """

    elif "High School" in level:
        return """
        The learner is in HIGH SCHOOL (Grade 10â€“12 in Kenyan CBC / KCSE level).
        âœ… Provide deeper understanding suitable for KCSE revision and critical thinking.
        âœ… Use more formal academic language and introduce subject terminology (define terms briefly).
        âœ… Explain logically with cause-effect reasoning or analytical breakdowns.
        âœ… Include exam-focused insights, marking scheme awareness, and common pitfalls.
        âœ… Use relatable teenage scenarios (career paths, innovation, social issues, science & tech).
        âœ… If explaining a concept:
            1. Provide a clear, exam-ready definition (1â€“2 lines).
            2. Break down the concept into 3â€“4 logical points or processes.
            3. Give one fully worked example/exercise with reasoning.
            4. Include an KCSE tip or misconception warning.
            5. Provide two exam-style practice questions (Standard + Advanced), with marking scheme guidance (steps and marks).
            6. End with a real-world, industry, or university-level application or why it matters for career growth.
        """

    else:  # College/University
        return """
        The learner is a COLLEGE/UNIVERSITY student (Year 1â€“5 or equivalent).
        âœ… Use advanced academic and technical language where appropriate.
        âœ… Provide structured and logically coherent explanations.
        âœ… Reference theories, frameworks, formulas, research insights, or case studies.
        âœ… Encourage critical thinking, comparison, and problem-solving.
        âœ… Tone should be professional yet clear and conceptually rich.
        âœ… If explaining a topic:
            1. Begin with a concise abstract/overview.
            2. Provide a detailed conceptual explanation with logical sectioning.
            3. Introduce formulas, models, or theoretical frameworks where relevant.
            4. Include a real-world or industry case study or application.
            5. Offer a solved advanced problem or scenario.
            6. Suggest further exploration or extension questions for deeper learning.
        """


if level=="-- Select Level --":
    st.warning("Please select your education level to proceed!")


mode = st.radio("Choose Study Mode:", ["ğŸ“„ Analyze Notes/Past Papers", "ğŸ’¬ Ask AI a Question","Mark My Answers"])
if mode == "ğŸ“„ Analyze Notes/Past Papers" and level=="-- Select Level --":
    st.warning("Please select your education level to proceed!")

elif mode == "ğŸ“„ Analyze Notes/Past Papers":
    lecture_file = st.file_uploader("ğŸ“˜ Upload Lecture Notes (PDF, TXT, or DOCX)", type=["pdf", "txt", "docx"])
    pastpaper_file = st.file_uploader("ğŸ“„ Upload Past Papers/Exams (PDF, TXT, or DOCX)", type=["pdf", "txt", "docx"])


    def extract_text(uploaded_file):
        if uploaded_file is None:
            return ""

        # Get file type
        file_type = uploaded_file.name.split(".")[-1].lower()

        # If it's a TXT file
        if file_type == "txt":
            return uploaded_file.read().decode("utf-8", errors="ignore")

        # If it's a PDF file
        elif file_type == "pdf":
            reader = PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text

        # If it's a DOCX file
        elif file_type == "docx":
            doc = Document(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text

        else:
            return "Unsupported file type"


    def extract_study_topics(lecture_text, pastpaper_text):
       prompt = f"""
       {get_level_prompt(level)}
       You are an educational AI assistant. Compare the following documents (lecture notes and past papers) and perform the following:
    
       1. Identify the **top 5 most frequently tested or emphasized concepts** based on past papers vs lecture content.
       2. For each concept, provide a **short, simple, and student-friendly explanation**.
       3. Be concise, clear, and avoid unnecessary jargon.
       4. Format the response using the structure below:
    
       **ğŸ“Œ KEY CONCEPTS:**
       - List the 5 concepts clearly in bullet form.
    
       **ğŸ“˜ EXPLANATIONS:**
       For each key concept, provide:
       Concept Name:
       Short Explanation (2â€“3 sentences max).
    
       Ensure the formatting is clean and easy for students to read and revise.
       Lecture Notes:
           {lecture_text[:]}
    
           Past Paper:
           {pastpaper_text[:]}
           """

       response = model.generate_content(prompt)
       return response.text

    def simplify(lecture_text, pastpaper_text):
        prompt = f"""
        {get_level_prompt(level)}
          You are a patient tutor who explains concepts in the simplest way possible using real-life analogies, examples, and step-by-step breakdowns. Assume the learner is a slow learner.

          Using the results below, explain each concept clearly in everyday language:

          RESULTS:
          {result}

          Now, based on the context in the lecture notes and past papers, further clarify using relatable analogies:

          LECTURE NOTES:
          {lecture_text}

          PAST PAPERS:
          {pastpaper_text}

          âœ… Your task:
          1. For each concept in the results, explain it as if teaching a slow learner.
          2. Use at least one everyday analogy for each concept.
          3. Break complex concepts into smaller steps.
          4. Give an easy example a high school student can understand.
          5. Keep explanations short, friendly, and encouraging.

          ğŸ“˜ Format like this:

          **Concept Name:**
          ğŸ”¹ Simple Explanation:
          ğŸ”¹ Analogy (real-life comparison):
          ğŸ”¹ Example:
          ğŸ”¹ Why it matters:

          Make it feel like a supportive tutor is guiding the student gently.
          """

        response = model.generate_content(prompt)
        return response.text

    def generate_practice_questions(lecture_text, pastpaper_text):
       prompt = f"""
                {get_level_prompt(level)}
                You are an expert academic examiner and educational AI. Carefully analyze and compare the concepts that appear in BOTH the lecture notes and past papers provided below. From the overlapping or recurring concepts:
        
               âœ… Generate exactly **30 well-structured, high-quality exam-style questions**.
               âœ… Use a natural mix of question types, such as:
                  - Short-answer questions
                  - Structured/descriptive questions
                  - Calculation or problem-solving questions (ONLY if applicable to the subject)
               âœ… Include a natural progression of difficulty (a blend of easier, moderately challenging, and advanced questions), but do NOT label or categorize difficulty levels.
               âœ… Ensure conceptual coverage is broad yet focused on repeated topics.
               âœ… Questions should feel professionally set, as in a formal college/university exam.
               âœ… DO NOT include multiple-choice questions.
               âœ… DO NOT provide any answers.
        
               ğŸ“˜ Format your response clearly as:
        
               **ğŸ“š EXAM QUESTION SET (30 Questions):**
        
               1. ...
               2. ...
               3. ...
               ...
               30. ...
        
               ---
        
               Here are the lecture notes:
               {lecture_text[:]}
        
               Here are the past papers:
               {pastpaper_text[:]}
             """

       response = model.generate_content(prompt)
       return response.text

    if lecture_file and pastpaper_file:
           with st.spinner("ğŸ¤– AI is thinking... hang tight!"):
               st.success("âœ…Files received")
               st.subheader("Iâ€™m now carefully analyzing your contentâ€”this may take a few seconds. â³")
               lecture_text = extract_text(lecture_file)
               pastpaper_text = extract_text(pastpaper_file)
               practice_questions = generate_practice_questions(lecture_text, pastpaper_text)
               with st.spinner("ğŸš€ Powering up your success... analyzing now!"):
                       result = extract_study_topics(lecture_text, pastpaper_text)
                       st.success("ğŸš€ Your AI-powered revision pack is ready!ğŸ”¥")
                       st.balloons()
                       st.subheader("Your personalized study guide is now ready - time to grow smarter")
                       st.write(result)

               if st.button("Simplify explanation"):
                           explanation=simplify(lecture_text, pastpaper_text)
                           st.write(explanation)
               st.download_button(
                           label="ğŸ“¥ Download Practice Questions",
                           data=practice_questions,
                           file_name="practice_questions.txt",
                           mime="text/plain"
                       )


    else:
        st.warning("Please upload both files in PDF, TXT, or DOCX format to continue.")
elif mode == "ğŸ’¬ Ask AI a Question" and level=="-- Select Level --":
    st.warning("Please select your education level to proceed!")

elif mode == "ğŸ’¬ Ask AI a Question":

    # Step 1: Ask for user input
    user_question = st.text_input("Ask a question:")

    # Step 2: Define the function (outside of button)
    def generate_answer(user_question):
        prompt = f"""
        {get_level_prompt(level)}
        You are an expert educational AI tutor who explains academic questions clearly and simply.
        The student asked:

        "{user_question}"

        Please respond with:
        ğŸ“˜ Explanation: Explain in a clear and simple way using an analogy if helpful.
        ğŸ“ Summary: Give a quick summary in 2 bullet points.
        ğŸ“ Practice Question: Provide 1 similar question for practice (without the answer).
        """

        response = model.generate_content(prompt)  # Your Gemini/OpenAI call
        return response.text  # Adjust if using OpenAI

    # Step 3: Show button and call function
    if st.button("Get Answer"):
        if user_question:
            answer = generate_answer(user_question)
            st.success(answer)
        else:
            st.warning("Please enter a question first.")
elif mode == "Mark My Answers" and level=="-- Select Level --":
    st.warning("Please select your education level to proceed!")

elif mode=="Mark My Answers":
    quiz = st.file_uploader("Upload your answered questions (PDF or TXT)", type=["pdf", "txt"])

    def answer_questions(question_file):
        prompt = f"""
            {get_level_prompt(level)}
            You are a certified examiner. I will give you:
            1. The answered exam by a student in pdf form. or Answered questions by a student in pdf form.
            Your task is to:
            âœ… Mark the student's answer objectively based on the exam standards.  
            âœ… Give a total score out of the maximum score in the question paper (default is 100 if not provided).  
            âœ… Highlight what the student did correctly and apllaud them.  
            âœ… Identify mistakes or missing key points.  
            âœ… Suggest improvements in simple language.  
            âœ… Provide a full model answer that would score 10/10.
            âš ï¸ IMPORTANT: Be strict but fair. Use marking scheme principles used in standard KCSE marking or College exams.
            Exam Question paper/exam:
                    {question_file}
            Now provide your result in the following structure:
            ğŸ¯ Score: X/maximum scrore possible then convert it to percentage form
            âœ… Correct Points:
            âŒ Mistakes / Missing Points:
            ğŸ“˜ Suggested Improvements:
            ğŸ“ Model Answer that would score all the marks (Perfect 10/10) or basically what would be the most most suitable answer to score the highest:
                        """
        response=model.generate_content(prompt)
        return response.text


    def generating_similar_questions(question_file):
        prompt = f"""
        {get_level_prompt(level)}
        You are an expert KCSE/WAEC question generator.
        Go through the answers provided by the student in {question_file}and identify every question the student got wrong.
        Your task is to:
        âœ… Understand the core concept being tested.  
        âœ… Generate new exam-style questions that test the same concept.
        âœ… Ensure each new question has the same difficulty level as the original.
        âœ… Do NOT provide answers unless I request them separately.
        âœ… Do NOT repeat or rephrase the original question.

        """
        response = model.generate_content(prompt)
        return response.text


    def extract_text(uploaded_file):
        if uploaded_file is None:
            return ""

        # Get file type
        file_type = uploaded_file.name.split(".")[-1].lower()

        # If it's a TXT file
        if file_type == "txt":
            return uploaded_file.read().decode("utf-8", errors="ignore")

        # If it's a PDF file
        elif file_type == "pdf":
            reader = PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text

        # If it's a DOCX file
        elif file_type == "docx":
            doc = Document(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text

        else:
            return "Unsupported file type"


    if quiz:
        st.subheader("ğŸ§  Letâ€™s see how you performed â€” AI is marking your paper...")
        with spinner("â³ Sit tight â€” AI is reviewing your answers. This may take a moment..."):
            question_file = extract_text(quiz)
            result=answer_questions(question_file)
            st.success("âœ… Marking complete!")
            st.subheader("ğŸ“Š Hereâ€™s your score and detailed feedback â€” letâ€™s see how you performed!")

            st.balloons()
            st.write(result)

            similar = generating_similar_questions(question_file)
            st.download_button(
                label="ğŸ“¥ Download Questions Similar To Those You Got Wrong ",
                data=similar,
                file_name="similar_questions.txt",
                mime="text/plain"
            )
    else:
        st.warning("Please upload your answered questions in PDF, TXT, or DOCX format to continue.")
st.markdown("""
<div style='text-align: center; padding-top: 30px; color: gray; font-size: 14px;'>
    ğŸ› ï¸ Built by <b>Papa YegoğŸ</b>
</div>
""", unsafe_allow_html=True)




